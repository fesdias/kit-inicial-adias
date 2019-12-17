#!/usr/bin/python
# -*- coding: utf-8 -*-

import datetime
import os
import re

from flask import Flask, render_template, request
from docx import Document
from pydrive.auth import GoogleAuth
from pydrive.drive import GoogleDrive

g_login = GoogleAuth()
g_login.LocalWebserverAuth()
drive = GoogleDrive(g_login)

app = Flask(__name__)

def dataAtual(doc):
    meses = {'1':"Janeiro", '2':"Fevereiro", '3':"Março", '4':"Abril", '5':"Maio", '6':"Junho", '7':"Julho", '8':"Agosto", '9':"Setembro", '10':"Outubro", '11':"Novembro", '12':"Dezembro"}
    data_atual = datetime.datetime.now()

    x = data_atual.strftime("%d")
    y = str(data_atual.strftime("%m"))

    if doc == "ADjudicia":
        z = int(data_atual.strftime("%Y"))+1
    else:
        z = data_atual.strftime("%Y")

    data = f"Jundiaí, {x} de {meses.get(y)} de {z}."

    return data

def preencher(doc, nomeDocumento, nome2, text, folderId):

    document = Document(doc)
    titulo = f"{nome2} - {nomeDocumento}.docx"
    doc = f"{nome2}/{titulo}"

    for paragraph in document.paragraphs:
        if "%data%" in paragraph.text:
            paragraph.text = dataAtual(nomeDocumento)

        if "%nome%" in paragraph.text:
            if nomeDocumento == "Contrato":
                paragraph.text = nome2.upper()
            else:
                paragraph.text = nome2

        for i in document.paragraphs:

            inline = i.runs
            for j in range(len(inline)):

                if inline[j].text == "cabecalhoaqui":
                    inline[j].text = text

                    try:
                        if inline[j-1].text == "%" or inline[j-1].text == "ü":
                            inline[j-1].text = ""

                        if inline[j+1].text == "%" or inline[j+1].text == "ü":
                            inline[j+1].text = ""

                    except:
                        pass


    document.save(doc)

    file = drive.CreateFile({'title': titulo, 'parents': [{'id': folderId}]})
    file.SetContentFile(doc)
    file.Upload()
    return titulo

def preencherTermo(nome, cpf, rg, end, cidade, cep, folderId):

    try:
        document = Document("Kitinicial/termo.docx")
    except:
        pass

    titulo = f"{nome} - Termo de Representação.docx"
    doc = f"{nome}/{titulo}"

    for paragraph in document.paragraphs:

        inline = paragraph.runs
        for j in range(len(inline)):

            if inline[j].text == "nomecompleto":
                inline[j].text = nome
                inline[j-1].text = ""
                inline[j+1].text = ""

            if inline[j].text == "numerocpf":
                inline[j].text = cpf
                inline[j-1].text = ""
                inline[j+1].text = ""

            if inline[j].text == "numerorg":
                inline[j].text = rg
                inline[j-1].text = ""
                inline[j+1].text = ""

            if inline[j].text == "enderecocompleto":
                inline[j].text = end
                inline[j-1].text = ""
                inline[j+1].text = ""

            if inline[j].text == "nomecidade":
                inline[j].text = cidade
                inline[j-1].text = ""
                inline[j+1].text = ""

            if inline[j].text == "numerocep":
                inline[j].text = cep
                inline[j-1].text = ""
                inline[j+1].text = ""

        document.save(doc)
        file = drive.CreateFile({'title': titulo, 'parents': [{'id': folderId}]})
        file.SetContentFile(doc)
        file.Upload()
        return titulo

@app.route('/')
def index():
    return render_template("Acervo.html")

@app.route("/enviar", methods = ["POST"])
def inserir():

    nome = request.form.get("nome")
    cpf = request.form.get("CPF")
    rg = request.form.get("RG")
    tel = request.form.get("tel")
    nacao = request.form.get("nacao")
    estCiv = request.form.get("estCiv")
    prof = request.form.get("prof")
    nasc = request.form.get("nasc")
    end = request.form.get("end")
    city = request.form.get("city")
    state = request.form.get("state")
    cep = request.form.get("cep")
    sexo = request.form.get("sexo")

    tipo = request.form.get("tipo")
    anda = request.form.get("andamento")
    status = request.form.get("status")

    nome = nome.upper()
    nome2 = nome.title()
    nacao = nacao.lower()
    estCiv = estCiv.lower()
    prof = prof.lower()
    end = end.title()
    city = city.title()

    if len(rg) == 8:
        rg = f"{rg[0]}{rg[1]}.{rg[2]}{rg[3]}{rg[4]}.{rg[5]}{rg[6]}{rg[7]}"

    if len(rg) == 9:
        rg = f"{rg[0]}{rg[1]}.{rg[2]}{rg[3]}{rg[4]}.{rg[5]}{rg[6]}{rg[7]}-{rg[8]}"

    if len(rg) == 10:
        rg = f"{rg[0]}{rg[1]}.{rg[2]}{rg[3]}{rg[4]}.{rg[5]}{rg[6]}{rg[7]}-{rg[8]}{rg[9]}"

    cpf = f"{cpf[0]}{cpf[1]}{cpf[2]}.{cpf[3]}{cpf[4]}{cpf[5]}.{cpf[6]}{cpf[7]}{cpf[8]}-{cpf[9]}{cpf[10]}"

    if sexo == 'f':
        text = f"{nome}, {nacao}, {estCiv}, {prof}, portadora do RG sob nº {rg}, e do CPF {cpf}, nascida em {nasc}, residente e domiciliada à {end}, {city}/{state} - CEP: {cep}"
    else:
        text = f"{nome}, {nacao}, {estCiv}, {prof}, portador do RG sob nº {rg}, e do CPF {cpf}, nascido em {nasc}, residente e domiciliado à {end}, {city}/{state} - CEP: {cep}"

    #Criar pasta sistema
    try:
        os.mkdir(nome2)
        pass
    except Exception as e:
        print("Já existe a pasta")
        pass

    #Criar pasta no drive
    folder = drive.CreateFile({'title': nome2, 'mimeType' : 'application/vnd.google-apps.folder', 'parents':[{'id':'1-1WHsERatmPkSH1dBgzyniANDWZdBRYQ'}]})
    folder.Upload()
    folderId = folder['id']
    #folderId = '1-1WHsERatmPkSH1dBgzyniANDWZdBRYQ'

    #Criar documento e salvar o caminho no sistema
    Proc = preencher("Kitinicial/procuracao.docx", "Procuração", nome2, text, folderId)
    Pobr = preencher("Kitinicial/pobreza.docx", "Pobreza", nome2, text, folderId)
    AD = preencher("Kitinicial/adjudicia.docx", "ADjudicia", nome2, text, folderId)
    Term = preencherTermo(nome2, cpf, rg, end, city, cep, folderId)

    if (sexo == "f"):
        Cont = preencher("Kitinicial/contrato-fem.docx", "Contrato", nome2, text, folderId)

    else:
        Cont = preencher("Kitinicial/contrato-masc.docx", "Contrato", nome2, text, folderId)

    path = os.getcwd()
    link = f"{path}/{nome2}"

    linkProc = f"{link}/{Proc}"
    linkProc = re.sub("\s", "%20", linkProc)

    linkPobr = f"{link}/{Pobr}"
    linkPobr = re.sub("\s", "%20", linkPobr)

    linkAD = f"{link}/{AD}"
    linkAD = re.sub("\s", "%20", linkAD)

    linkTermo = f"{link}/{Term}"
    linkTermo = re.sub("\s", "%20", linkTermo)

    linkCont = f"{link}/{Cont}"
    linkCont = re.sub("\s", "%20", linkCont)

    concluido = f"Encontre os documentos de {nome2} em"

    return render_template("confirmacao.html", text = concluido, cab = text, linkProc = linkProc, linkPobr = linkPobr, linkAD = linkAD, linkTermo = linkTermo)

@app.route("/buscar", methods = ["POST", "GET"])
def buscar():

    buscar = request.form.get("buscar")

    clientes = {}
    return render_template("naoencontrado.html")
